import os
import json
from pathlib import Path
import sys
import zipfile
import xml.etree.ElementTree as ET
from docxtpl import DocxTemplate, RichText
from docx import Document
from docx2pdf import convert
import fitz
import piexif
from shapely.geometry import LineString, Point
from shapely.ops import transform
from pyproj import Transformer
import subprocess
from datetime import datetime
import shutil
import logging
import re
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "LOG.txt")
logging.basicConfig(filename=LOG_FILE, level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

try:
    with open(Path(__file__).parent / "Pilot Directory" / "pilot_names.json", "r") as f:
        MASTER_PILOTS = json.load(f)["pilots"]
except Exception as e:
    logging.error("Failed to load pilot_names.json: %s", e)
    MASTER_PILOTS = []
_transformer = Transformer.from_crs("epsg:4326", "epsg:3857", always_xy=True)
KML_NS = {"kml": "http://www.opengis.net/kml/2.2"}
UNC_FOLDER = "Uncompressed Reports"
FOLDERS_TO_MOVE = ["Reports by Day", "Summary Report", "Report by Pipeline"]
MOVE_JPGS = True
OUTPUT_FOLDER_NAME = "__CLIENT DELIVERABLES__"
MAIN_REPORTS_SUBFOLDER = "Main Reports"
PIPELINE_REPORTS_SUBFOLDER = "Reports by Pipeline"
KEYWORDS_FOR_MAIN_REPORTS = ["Summary", "Encroachment"]
ORIGINAL_PHOTOS_FOLDER = "ORIGINAL_PHOTOS"

PILOT_NOTE_MAP = {
    "C_Active": "Active Construction near ROW",
    "C_Inactive": "Inactive Construction near ROW",
    "E_Active": "Active Excavation near ROW",
    "E_Inactive": "Inactive Excavation near ROW",
    "F_Pipeline": "Exposed Pipeline near ROW",
    "O_Equipment": "Equipment near ROW",
    "O_Sheen": "Sheening in water near ROW",
    "O_Landscape": "Impacted Landscape near ROW",
    "O_Powerline": "Powerline Operations near ROW",
    "O_Sign": "Impacted sign near ROW",
    "O_Railroad": "Railroad Operations near ROW",
    "O_Boom": "Boom in water near ROW",
    "O_Trash": "Trash near ROW"
}

def create_kml(lat, lon, name, img_filename):
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <Placemark>
      <name>{name}</name>
      <description><![CDATA[<img src="{img_filename}" width="300"/>]]></description>
      <Point>
        <coordinates>{lon},{lat},0</coordinates>
      </Point>
    </Placemark>
  </Document>
</kml>
"""

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(underline)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(rPr)
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return paragraph

def clone_row(table, row):
    tbl = table._tbl
    new_tr = copy.deepcopy(row._tr)
    tbl.append(new_tr)
    return table.rows[-1]

def fix_header_pilot_names(doc):
    """Format pilot names in the document header without removing other content."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    header = doc.sections[0].header

    # locate the paragraph that introduces pilot names
    paragraphs = list(header.paragraphs)
    label_idx = None
    for i, para in enumerate(paragraphs):
        if para.text.strip().lower().startswith("pilot names"):
            label_idx = i
            break

    if label_idx is None:
        return

    # gather pilot name lines following the label
    names = []
    for para in paragraphs[label_idx + 1:]:
        text = para.text.strip()
        if text in ("{% for p in PILOT_NAMES %}", "{% endfor %}", "{{ p }}"):
            para._element.getparent().remove(para._element)
            continue
        if text:
            if text.startswith("- "):
                text = text[2:].strip()
            names.append(text)
        para._element.getparent().remove(para._element)

    label_para = header.paragraphs[label_idx]
    base_indent = label_para.paragraph_format.left_indent
    for name in names:
        p = header.add_paragraph(name, style="Header")
        label_para._element.addnext(p._element)
        label_para = p
        for run in p.runs:
            run.font.name = "Aptos Display"
            run.font.size = Pt(10)
            run.font.bold = False
            run.font.underline = False
            run.font.italic = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = base_indent
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

class ClientConfig:
    def __init__(self):
        self.base_dir = Path(__file__).parent

    def get_settings(self, client):
        cfg_file = self.base_dir / "Clients" / client / "config.json"
        with cfg_file.open("r") as f:
            settings = json.load(f)
        client_dir = self.base_dir / "Clients" / client
        templates_dir = client_dir / "Templates - DOCX"
        if templates_dir.is_dir():
            if not settings.get("template_pipeline"):
                pipeline_files = [f for f in os.listdir(templates_dir) if f.lower().endswith('.docx') and "pipeline" in f.lower()]
                if pipeline_files:
                    settings["template_pipeline"] = str(templates_dir / pipeline_files[0])
            if not settings.get("template_summary"):
                summary_files = [f for f in os.listdir(templates_dir) if f.lower().endswith('.docx') and "summary" in f.lower()]
                if summary_files:
                    settings["template_summary"] = str(templates_dir / summary_files[0])
        # normalize template paths to absolute
        for key in ["template_pipeline", "template_summary"]:
            if settings.get(key):
                p = Path(settings[key])
                if not p.is_absolute():
                    settings[key] = str((self.base_dir / p).resolve())
        kmz_dir = client_dir / "Pipeline Systems - KMZ"
        if kmz_dir.is_dir():
            kmz_files = [f for f in os.listdir(kmz_dir) if f.lower().endswith('.kmz')]
            settings["kmz_path"] = str(kmz_dir / kmz_files[0]) if kmz_files else ""
        settings["report_prefix"] = settings.get("report_prefix", client)
        return settings

class KMZParser:
    def __init__(self, client="HGS"):
        self.client = client

    def _sanitize(self, name):
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in name.strip())

    def _extract_kml_from_kmz(self, kmz_path):
        with zipfile.ZipFile(kmz_path) as z:
            for name in z.namelist():
                if name.lower().endswith(".kml"):
                    with z.open(name) as f:
                        return f.read()
        raise RuntimeError("No KML in KMZ")

    def load_pipelines(self, kmz_path):
        kml_data = self._extract_kml_from_kmz(kmz_path)
        root = ET.fromstring(kml_data)
        pipelines = []
        for placemark in root.findall('.//kml:Placemark', KML_NS):
            ls = placemark.find('.//kml:LineString', KML_NS)
            if ls is None:
                continue
            name_el = placemark.find('kml:name', KML_NS)
            name = name_el.text.strip() if name_el is not None else 'pipeline'
            coord_el = ls.find('kml:coordinates', KML_NS)
            if coord_el is None or not coord_el.text:
                continue
            coords = []
            for part in coord_el.text.strip().split():
                pieces = part.split(',')
                if len(pieces) >= 2:
                    lon, lat = map(float, pieces[:2])
                    coords.append((lon, lat))
            if len(coords) >= 2:
                pipelines.append((name, LineString(coords)))
        return pipelines

    def rename_images(self, folder):
        config = ClientConfig()
        settings = config.get_settings(self.client)
        kmz_path = settings.get("kmz_path")
        if not kmz_path or not os.path.exists(kmz_path):
            logging.error(f"No valid KMZ path for {self.client}")
            return
        pipelines = self.load_pipelines(kmz_path)
        if not pipelines:
            logging.error("No pipelines loaded")
            return
        buffered = [(self._sanitize(name), transform(_transformer.transform, line).buffer(15.24)) for name, line in pipelines]
        kml_folder = os.path.join(folder, "KMLs")
        os.makedirs(kml_folder, exist_ok=True)
        images = [f for f in os.listdir(folder) if f.lower().endswith(('.jpg', '.jpeg'))]
        for img in images:
            img_path = os.path.join(folder, img)
            coords = self._read_photo_coords(img_path)
            if not coords:
                logging.warning(f"No GPS data for {img}")
                continue
            x, y = _transformer.transform(coords[1], coords[0])
            point = Point(x, y)
            hit_names = [name for name, buf in buffered if buf.contains(point)]
            if not hit_names:
                logging.info(f"{img} did not match any pipeline")
                continue

            # Check if client abbreviation is already in the image name
            img_base = os.path.splitext(img)[0]
            if f"_{self.client}_" in f"_{img_base}_":  # Ensure exact match with underscores
                # Client abbreviation present, don't add it again
                if len(hit_names) == 1:
                    combined_name = f"{hit_names[0]}_{img}"
                    new_name = re.sub(r'__+', '_', combined_name)
                    new_path = os.path.join(folder, new_name)
                    os.rename(img_path, new_path)
                    lat, lon = coords
                    kml_name = os.path.splitext(new_name)[0]
                    kml_content = create_kml(lat, lon, kml_name, new_name)
                    with open(os.path.join(kml_folder, f"{kml_name}.kml"), "w", encoding="utf-8") as f:
                        f.write(kml_content)
                else:
                    for name in hit_names:
                        combined_name = f"{name}_{img}"
                        new_name = re.sub(r'__+', '_', combined_name)
                        new_path = os.path.join(folder, new_name)
                        shutil.copy2(img_path, new_path)
                        lat, lon = coords
                        kml_name = os.path.splitext(new_name)[0]
                        kml_content = create_kml(lat, lon, kml_name, new_name)
                        with open(os.path.join(kml_folder, f"{kml_name}.kml"), "w", encoding="utf-8") as f:
                            f.write(kml_content)
                    os.remove(img_path)
            else:
                # Client abbreviation not present, add it
                if len(hit_names) == 1:
                    combined_name = f"{hit_names[0]}_{self.client}_{img}"
                    new_name = re.sub(r'__+', '_', combined_name)
                    new_path = os.path.join(folder, new_name)
                    os.rename(img_path, new_path)
                    lat, lon = coords
                    kml_name = os.path.splitext(new_name)[0]
                    kml_content = create_kml(lat, lon, kml_name, new_name)
                    with open(os.path.join(kml_folder, f"{kml_name}.kml"), "w", encoding="utf-8") as f:
                        f.write(kml_content)
                else:
                    for name in hit_names:
                        combined_name = f"{name}_{self.client}_{img}"
                        new_name = re.sub(r'__+', '_', combined_name)
                        new_path = os.path.join(folder, new_name)
                        shutil.copy2(img_path, new_path)
                        lat, lon = coords
                        kml_name = os.path.splitext(new_name)[0]
                        kml_content = create_kml(lat, lon, kml_name, new_name)
                        with open(os.path.join(kml_folder, f"{kml_name}.kml"), "w", encoding="utf-8") as f:
                            f.write(kml_content)
                    os.remove(img_path)

    def _read_photo_coords(self, photo_path):
        try:
            exif_data = piexif.load(photo_path)
            gps = exif_data.get("GPS", {})
            lat = gps.get(piexif.GPSIFD.GPSLatitude)
            lon = gps.get(piexif.GPSIFD.GPSLongitude)
            lat_ref = gps.get(piexif.GPSIFD.GPSLatitudeRef)
            lon_ref = gps.get(piexif.GPSIFD.GPSLongitudeRef)
            if None in (lat, lon, lat_ref, lon_ref):
                return None
            def _to_deg(value):
                d, m, s = value
                return d[0]/d[1] + m[0]/m[1]/60 + s[0]/s[1]/3600
            lat_deg = _to_deg(lat)
            lon_deg = _to_deg(lon)
            if lat_ref in [b'S', b's']:
                lat_deg = -lat_deg
            if lon_ref in [b'W', b'w']:
                lon_deg = -lon_deg
            return lat_deg, lon_deg
        except Exception as e:
            logging.error(f"EXIF error in {photo_path}: {e}")
            return None

class ReportGenerator:
    IMAGE_BOX = fitz.Rect(49.64, 78.52, 576.64, 375.35)
    ROW_HEIGHT_TWIPS = '648'

    def __init__(self, config):
        self.config = config

    def extract_metadata_for_pipeline_report(self, image_path):
        metadata = {
            "FILE_NAME": os.path.basename(image_path),
            "PLACEHOLDER_DATE_TIME": "Unavailable",
            "LATITUDE": "Unavailable",
            "LONGITUDE": "Unavailable",
            "GOOGLE_URL": "Unavailable",
            "PILOT_NOTES": "",
            "PIPELINE_ID": "Unavailable"
        }
        filename = os.path.basename(image_path)
        base = filename.rsplit(".", 1)[0]
        parts = base.split("_")
        try:
            client_index = parts.index(self.client)
            if client_index > 0:
                metadata["PIPELINE_ID"] = "_".join(parts[:client_index]).replace("_", " ")
        except ValueError:
            pass
        try:
            exif_dict = piexif.load(image_path)
            date_bytes = exif_dict["Exif"].get(piexif.ExifIFD.DateTimeOriginal)
            if date_bytes:
                metadata["PLACEHOLDER_DATE_TIME"] = datetime.strptime(date_bytes.decode("utf-8"), "%Y:%m:%d %H:%M:%S").strftime("%Y-%m-%d %H:%M:%S")
            gps = exif_dict.get("GPS", {})
            if piexif.GPSIFD.GPSLatitude in gps and piexif.GPSIFD.GPSLongitude in gps:
                lat = self._to_decimal(gps[piexif.GPSIFD.GPSLatitude], gps[piexif.GPSIFD.GPSLatitudeRef])
                lon = self._to_decimal(gps[piexif.GPSIFD.GPSLongitude], gps[piexif.GPSIFD.GPSLongitudeRef])
                metadata["LATITUDE"] = f"{lat:.6f}"
                metadata["LONGITUDE"] = f"{lon:.6f}"
                metadata["GOOGLE_URL"] = f"https://maps.google.com/?q={lat:.6f},{lon:.6f}"
        except Exception as e:
            logging.error(f"EXIF error in {image_path}: {e}")
        return metadata

    def _to_decimal(self, dms, ref):
        d, m, s = dms
        deg = d[0]/d[1] + m[0]/m[1]/60 + s[0]/s[1]/3600
        return -deg if ref in [b'S', b'W'] else deg

    def extract_metadata_for_summary_report(self, image_path):
        metadata = {
            "PIPELINE_ID": "Unavailable",
            "Photo_ID": "Unavailable",
            "LAT_LONG": "Unavailable",
            "GOOGLE_URL": "Unavailable",
            "CODE": "Unavailable",
            "STATUS": "Unavailable",
            "NOTE_FIELD": ""
        }
        filename = os.path.basename(image_path).rsplit(".", 1)[0]
        parts = filename.split("_")
        try:
            client_index = parts.index(self.client)
            if client_index is not None and len(parts) > client_index + 2:
                metadata["PIPELINE_ID"] = "_".join(parts[:client_index])
                metadata["Photo_ID"] = parts[client_index] + "_" + parts[client_index + 1]
                metadata["CODE"] = parts[client_index + 2] if len(parts) > client_index + 2 else ""
                metadata["STATUS"] = parts[client_index + 3] if len(parts) > client_index + 3 else ""
                key = f"{metadata['CODE']}_{metadata['STATUS'].replace(' ', '')}"
                metadata["NOTE_FIELD"] = PILOT_NOTE_MAP.get(key, "")
        except ValueError:
            metadata["Photo_ID"] = filename
        try:
            exif_dict = piexif.load(image_path)
            gps = exif_dict.get("GPS", {})
            if piexif.GPSIFD.GPSLatitude in gps and piexif.GPSIFD.GPSLongitude in gps:
                lat = self._to_decimal(gps[piexif.GPSIFD.GPSLatitude], gps[piexif.GPSIFD.GPSLatitudeRef])
                lon = self._to_decimal(gps[piexif.GPSIFD.GPSLongitude], gps[piexif.GPSIFD.GPSLongitudeRef])
                metadata["LAT_LONG"] = f"({lat:.6f}, {lon:.6f})"
                metadata["GOOGLE_URL"] = f"https://maps.google.com/?q={lat:.6f},{lon:.6f}"
        except Exception as e:
            logging.error(f"EXIF error in {image_path}: {e}")
        return metadata

    def generate_reports(self, folder, images, client="HGS", cover_docx=None, *, pilot_names=None):
        self.client = client
        pilot_names = pilot_names or []
        settings = self.config.get_settings(client)
        if not settings.get("template_pipeline") or not settings.get("template_summary"):
            logging.error(f"Templates not found for {client}")
            return
        current_date = datetime.now().strftime("%Y_%m_%d")

        # Pipeline Reports
        report_dir = os.path.join(folder, "Report by Pipeline")
        os.makedirs(report_dir, exist_ok=True)
        for img in images:
            img_path = os.path.join(folder, img)
            context = self.extract_metadata_for_pipeline_report(img_path)
            def link(text, url):
                return RichText(url, color="0000FF", underline=True)
            context["link"] = link
            context["PILOT_NAMES"] = pilot_names
            image_name_no_ext = os.path.splitext(img)[0]
            image_report_folder = os.path.join(report_dir, image_name_no_ext)
            os.makedirs(image_report_folder, exist_ok=True)
            docx_filename = f"CONFIDENTIAL_AALLC_{image_name_no_ext}_{current_date}.docx"
            pdf_filename = f"CONFIDENTIAL_AALLC_{image_name_no_ext}_{current_date}.pdf"
            docx_output_path = os.path.join(image_report_folder, docx_filename)
            pdf_output_path = os.path.join(image_report_folder, pdf_filename)
            temp_pdf_path = os.path.join(image_report_folder, f"_temp_{image_name_no_ext}.pdf")
            tpl = DocxTemplate(settings["template_pipeline"])
            tpl.render(context)
            tpl.save(docx_output_path)
            convert(docx_output_path, temp_pdf_path)
            doc = fitz.open(temp_pdf_path)
            if len(doc) > 0:
                doc[0].insert_image(self.IMAGE_BOX, filename=img_path, keep_proportion=False)
                doc.save(pdf_output_path)
                doc.close()
                os.remove(temp_pdf_path)
            else:
                os.rename(temp_pdf_path, pdf_output_path)

        # Day Reports
        reports_dir = os.path.join(folder, "Reports by Day")
        os.makedirs(reports_dir, exist_ok=True)
        pdf_files = [os.path.join(report_dir, os.path.splitext(img)[0], f"CONFIDENTIAL_AALLC_{os.path.splitext(img)[0]}_{current_date}.pdf") for img in images]
        output_path = os.path.join(reports_dir, f"CONFIDENTIAL_AALLC_{settings['report_prefix']}_EncroachmentPatrols_{current_date}_Report.pdf")
        cover_pdf_path = None
        if cover_docx and os.path.exists(cover_docx):
            cover_pdf_path = os.path.join(reports_dir, "_cover_temp.pdf")
            convert(cover_docx, cover_pdf_path)

        merged_doc = fitz.open()
        if cover_pdf_path and os.path.exists(cover_pdf_path):
            cover_doc = fitz.open(cover_pdf_path)
            merged_doc.insert_pdf(cover_doc)
            cover_doc.close()
        for pdf_path in pdf_files:
            if os.path.exists(pdf_path):
                pdf_doc = fitz.open(pdf_path)
                merged_doc.insert_pdf(pdf_doc)
                pdf_doc.close()
        merged_doc.save(output_path)
        merged_doc.close()
        if cover_pdf_path and os.path.exists(cover_pdf_path):
            os.remove(cover_pdf_path)

        # Summary Report
        summary_dir = os.path.join(folder, "Summary Report")
        os.makedirs(summary_dir, exist_ok=True)
        tpl = DocxTemplate(settings["template_summary"])
        tpl.render({"PILOT_NAMES": pilot_names})
        temp_summary_tpl = os.path.join(summary_dir, "_summary_temp.docx")
        tpl.save(temp_summary_tpl)

        doc = Document(temp_summary_tpl)
        fix_header_pilot_names(doc)
        table = doc.tables[2]
        placeholder = next(
            (row for row in table.rows if any("{{" in cell.text for cell in row.cells)),
            None,
        )
        if placeholder is None:
            if len(table.rows) >= 2:
                placeholder = table.rows[1]
            else:
                placeholder = table.rows[0]

        for img in images:
            img_path = os.path.join(folder, img)
            metadata = self.extract_metadata_for_summary_report(img_path)
            new_row = clone_row(table, placeholder)
            tr = new_row._tr
            trPr = tr.get_or_add_trPr()
            for existing in trPr.findall(qn("w:trHeight")):
                trPr.remove(existing)
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), self.ROW_HEIGHT_TWIPS)
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)
            for idx, key in enumerate(["PIPELINE_ID", "Photo_ID", "LAT_LONG", "CODE", "STATUS", "NOTE_FIELD"]):
                if idx >= len(new_row.cells):
                    continue
                cell = new_row.cells[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if key == "LAT_LONG":
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    add_hyperlink(paragraph, metadata["GOOGLE_URL"], metadata["LAT_LONG"])
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell.text = metadata.get(key, "")
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table._tbl.remove(placeholder._tr)
        out_path = os.path.join(summary_dir, f"CONFIDENTIAL_AALLC_{settings['report_prefix']}_{current_date}_PatrolSummaryReport.docx")
        pdf_path = os.path.join(summary_dir, f"CONFIDENTIAL_AALLC_{settings['report_prefix']}_{current_date}_PatrolSummaryReport.pdf")
        doc.save(out_path)
        convert(out_path, pdf_path)
        if os.path.exists(temp_summary_tpl):
            os.remove(temp_summary_tpl)

class PDFProcessor:
    def find_pdfs(self, root_folder):
        pdfs = []
        for dirpath, _, filenames in os.walk(root_folder):
            for f in filenames:
                if f.lower().endswith('.pdf'):
                    pdfs.append(os.path.join(dirpath, f))
        return pdfs

    def compress_with_ghostscript(self, infile, outfile):
        gs_path = "gswin64c" if sys.platform == "win32" else "gs"
        args = [
            gs_path, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.6",
            "-dPDFSETTINGS=/ebook", "-dNOPAUSE", "-dQUIET", "-dBATCH",
            f"-sOutputFile={outfile}", infile
        ]
        try:
            subprocess.run(args, check=True)
            return True
        except Exception as e:
            logging.error(f"Compression failed for {infile}: {e}")
            return False

    def safe_filename(self, filename, used_names):
        base, ext = os.path.splitext(filename)
        i = 1
        newname = filename
        while newname in used_names:
            newname = f"{base}_{i}{ext}"
            i += 1
        used_names.add(newname)
        return newname

    def compress_pdfs(self, folder):
        pdf_files = self.find_pdfs(folder)
        if not pdf_files:
            logging.info("No PDFs found!")
            return
        output_dir = os.path.join(folder, OUTPUT_FOLDER_NAME)
        os.makedirs(output_dir, exist_ok=True)
        used_names = set()
        for pdf_path in pdf_files:
            original_name = os.path.basename(pdf_path)
            out_name = self.safe_filename(original_name, used_names)
            out_path = os.path.join(output_dir, out_name)
            self.compress_with_ghostscript(pdf_path, out_path)

        main_reports_dir = os.path.join(output_dir, MAIN_REPORTS_SUBFOLDER)
        os.makedirs(main_reports_dir, exist_ok=True)
        for filename in os.listdir(output_dir):
            if not filename.lower().endswith(".pdf"):
                continue
            if any(keyword in filename for keyword in KEYWORDS_FOR_MAIN_REPORTS):
                src = os.path.join(output_dir, filename)
                dst = os.path.join(main_reports_dir, filename)
                if os.path.exists(src):
                    shutil.move(src, dst)

        pipeline_reports_dir = os.path.join(output_dir, PIPELINE_REPORTS_SUBFOLDER)
        os.makedirs(pipeline_reports_dir, exist_ok=True)
        for filename in os.listdir(output_dir):
            if not filename.lower().endswith(".pdf"):
                continue
            src = os.path.join(output_dir, filename)
            dst = os.path.join(pipeline_reports_dir, filename)
            if os.path.exists(src):
                shutil.move(src, dst)

        uncompressed_dir = os.path.join(folder, UNC_FOLDER)
        os.makedirs(uncompressed_dir, exist_ok=True)
        for foldername in FOLDERS_TO_MOVE:
            src_path = os.path.join(folder, foldername)
            dst_path = os.path.join(uncompressed_dir, foldername)
            if os.path.isdir(src_path):
                shutil.move(src_path, dst_path)
        if MOVE_JPGS:
            for fname in os.listdir(folder):
                if fname.lower().endswith((".jpg", ".jpeg")) and os.path.isfile(os.path.join(folder, fname)):
                    src = os.path.join(folder, fname)
                    dst = os.path.join(uncompressed_dir, fname)
                    shutil.move(src, dst)
        src_photos = os.path.join(folder, ORIGINAL_PHOTOS_FOLDER)
        dst_photos = os.path.join(output_dir, ORIGINAL_PHOTOS_FOLDER)
        if os.path.isdir(src_photos):
            shutil.move(src_photos, dst_photos)

class ReportApp:
    """Launcher class that binds the CustomTkinter GUI to report logic."""

    def __init__(self):
        base = Path(__file__).parent
        self.clients = [d.name for d in (base / "Clients").iterdir() if (d / "config.json").exists()]
        if len(sys.argv) > 1 and os.path.isdir(sys.argv[1]):
            self.initial_folder = sys.argv[1]
        else:
            self.initial_folder = None

    def run(self):
        from custom_ui import ReportGUI

        def callback(folder, client, cover, pilots, set_progress, set_status):
            self._generate(folder, client, cover, pilots, set_progress, set_status)

        logo_file = Path(__file__).parent / "Arch_Aerial_LOGO.jpg"
        gui = ReportGUI(
            MASTER_PILOTS,
            self.clients,
            callback,
            logo_path=logo_file,
            initial_folder=self.initial_folder,
        )
        gui.mainloop()

    def _generate(self, folder, client, cover, pilots, set_progress, set_status):
        if not os.path.isdir(folder):
            set_status(text="Invalid folder")
            return

        set_status(text="Backing up photos...")
        orig_folder = os.path.join(folder, ORIGINAL_PHOTOS_FOLDER)
        os.makedirs(orig_folder, exist_ok=True)
        images = [f for f in os.listdir(folder) if f.lower().endswith((".jpg", ".jpeg"))]
        for img in images:
            shutil.copy2(os.path.join(folder, img), os.path.join(orig_folder, img))
        set_progress(0.25)

        set_status(text="Renaming images...")
        kmz = KMZParser(client)
        kmz.rename_images(folder)
        images = [f for f in os.listdir(folder) if f.lower().endswith((".jpg", ".jpeg"))]
        set_progress(0.5)

        set_status(text="Generating reports...")
        config = ClientConfig()
        report_gen = ReportGenerator(config)
        cover_doc = cover if cover and os.path.exists(cover) else None
        report_gen.generate_reports(folder, images, client, cover_doc, pilot_names=pilots)
        set_progress(0.75)

        set_status(text="Compressing PDFs...")
        pdf_proc = PDFProcessor()
        pdf_proc.compress_pdfs(folder)
        set_progress(1.0)

        set_status(text="Reports done!")

if __name__ == "__main__":
    app = ReportApp()
    app.run()